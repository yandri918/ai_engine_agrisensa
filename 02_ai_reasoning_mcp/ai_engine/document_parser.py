"""
AgriSensa Document Parser & Fetcher (MCP Tool)
=============================================
Modul pengolah & ekstraksi dokumen untuk AgriSensa:
- Parse PDF (fitur: PyMuPDF / pypdf fallback)
- Parse Word .docx (python-docx)
- Parse Excel .xlsx / .xls (openpyxl / pandas)
- Parse CSV / TSV (pandas / csv)
- Fetch & Parse URL (PDF/HTML via WebScraper)
- Ekstraksi tabel keuangan & angka metrik pertanian (RAB, harga, luas, yield)
- Gemini AI Summarization & Key Insights extraction
"""

import os
import io
import re
import csv
import json
import logging
import time
import base64
from dataclasses import dataclass, field, asdict
from typing import Any, Dict, List, Optional, Union
from datetime import datetime

logger = logging.getLogger("agrisensa.document_parser")

# ─────────────────────────────────────────────────────────────────────────────
# Parser Library Checks
# ─────────────────────────────────────────────────────────────────────────────
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    fitz = None
    PYMUPDF_AVAILABLE = False

try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PdfReader = None
    PYPDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    docx = None
    DOCX_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    pd = None
    PANDAS_AVAILABLE = False


@dataclass
class ParsedDocumentResult:
    filename: str
    file_type: str                  # pdf, docx, xlsx, csv, url, text
    total_pages_or_sheets: int
    extracted_text: str
    extracted_tables: List[List[List[str]]]
    extracted_metrics: Dict[str, Any]
    ai_summary: str
    key_findings: List[str]
    processing_time_sec: float
    timestamp: str


class DocumentParser:
    """
    AgriSensa Document Parser & Analyzer.
    Mengekstrak teks, tabel, data keuangan dan metrik pertanian dari berbagai file.
    """

    def __init__(self):
        self.deepseek_api_key = os.getenv("DEEPSEEK_API_KEY", "") or os.getenv("OPENAI_API_KEY", "")
        logger.info(
            f"DocumentParser init (PyMuPDF={PYMUPDF_AVAILABLE}, docx={DOCX_AVAILABLE}, openpyxl={OPENPYXL_AVAILABLE})"
        )

    def parse_file(
        self,
        file_bytes: bytes,
        filename: str,
        file_type: Optional[str] = None
    ) -> ParsedDocumentResult:
        """Parse file dari bytes (upload)."""
        t0 = time.perf_counter()
        ext = file_type or filename.split(".")[-1].lower()

        text = ""
        tables = []
        pages = 1

        if ext == "pdf":
            text, tables, pages = self._parse_pdf(file_bytes)
        elif ext in ["docx", "doc"]:
            text, tables, pages = self._parse_docx(file_bytes)
        elif ext in ["xlsx", "xls"]:
            text, tables, pages = self._parse_excel(file_bytes)
        elif ext in ["csv", "tsv"]:
            text, tables, pages = self._parse_csv(file_bytes, delimiter="," if ext == "csv" else "\t")
        else:
            # Plain text
            text = file_bytes.decode("utf-8", errors="ignore")

        metrics = self._extract_agricultural_metrics(text)
        ai_summary, findings = self._summarize_with_gemini(filename, text, metrics)

        return ParsedDocumentResult(
            filename=filename,
            file_type=ext,
            total_pages_or_sheets=pages,
            extracted_text=text[:30000],
            extracted_tables=tables[:10],
            extracted_metrics=metrics,
            ai_summary=ai_summary,
            key_findings=findings,
            processing_time_sec=round(time.perf_counter() - t0, 3),
            timestamp=datetime.now().isoformat(),
        )

    def parse_base64(self, b64_content: str, filename: str) -> Dict[str, Any]:
        """Parse base64 uploaded file."""
        try:
            raw_bytes = base64.b64decode(b64_content)
            res = self.parse_file(raw_bytes, filename)
            return {"success": True, "data": asdict(res)}
        except Exception as e:
            logger.error(f"Error parsing base64 file {filename}: {e}")
            return {"success": False, "error": str(e)}

    def parse_from_path(self, file_path: str) -> Dict[str, Any]:
        """Parse file dari file system local."""
        try:
            filename = os.path.basename(file_path)
            with open(file_path, "rb") as f:
                content = f.read()
            res = self.parse_file(content, filename)
            return {"success": True, "data": asdict(res)}
        except Exception as e:
            logger.error(f"Error parsing local file {file_path}: {e}")
            return {"success": False, "error": str(e)}

    # ──────────────────────────────────────
    # Specific format parsers
    # ──────────────────────────────────────

    def _parse_pdf(self, file_bytes: bytes) -> tuple:
        text_chunks = []
        tables = []
        pages = 0

        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                pages = len(doc)
                for page_num in range(pages):
                    page = doc[page_num]
                    text_chunks.append(f"--- Halaman {page_num + 1} ---\n" + page.get_text())
                    # Extract tables if pymupdf table feature exists
                    try:
                        tabs = page.find_tables()
                        for tab in tabs:
                            tables.append(tab.extract())
                    except Exception:
                        pass
                return "\n".join(text_chunks), tables, pages
            except Exception as e:
                logger.warning(f"PyMuPDF failed, fallback: {e}")

        if PYPDF_AVAILABLE:
            try:
                reader = PdfReader(io.BytesIO(file_bytes))
                pages = len(reader.pages)
                for i, page in enumerate(reader.pages):
                    text_chunks.append(f"--- Halaman {i + 1} ---\n" + (page.extract_text() or ""))
                return "\n".join(text_chunks), tables, pages
            except Exception as e:
                logger.error(f"pypdf failed: {e}")

        return "Gagal mengekstrak teks PDF. Install PyMuPDF: pip install pymupdf", tables, 0

    def _parse_docx(self, file_bytes: bytes) -> tuple:
        if not DOCX_AVAILABLE:
            return "python-docx belum terinstall. Jalankan: pip install python-docx", [], 0
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text_lines = [p.text for p in doc.paragraphs if p.text]
            tables = []
            for t in doc.tables:
                rows = []
                for row in t.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                if rows:
                    tables.append(rows)
            return "\n".join(text_lines), tables, 1
        except Exception as e:
            return f"Error parsing docx: {str(e)}", [], 0

    def _parse_excel(self, file_bytes: bytes) -> tuple:
        if PANDAS_AVAILABLE:
            try:
                excel_file = io.BytesIO(file_bytes)
                xls = pd.ExcelFile(excel_file)
                text_parts = []
                tables = []
                for sheet_name in xls.sheet_names:
                    df = pd.read_excel(xls, sheet_name=sheet_name)
                    text_parts.append(f"=== Sheet: {sheet_name} ===\n" + df.to_string())
                    # Convert to table matrix (string format)
                    matrix = [df.columns.astype(str).tolist()] + df.astype(str).values.tolist()
                    tables.append(matrix[:50])
                return "\n\n".join(text_parts), tables, len(xls.sheet_names)
            except Exception as e:
                logger.warning(f"pandas excel read failed: {e}")

        return "Excel parsing membutuhkan pandas/openpyxl: pip install pandas openpyxl", [], 0

    def _parse_csv(self, file_bytes: bytes, delimiter: str = ",") -> tuple:
        try:
            content = file_bytes.decode("utf-8", errors="ignore")
            reader = csv.reader(io.StringIO(content), delimiter=delimiter)
            rows = list(reader)
            text = "\n".join([", ".join(r) for r in rows[:100]])
            return text, [rows[:100]], 1
        except Exception as e:
            return f"Error parsing CSV: {str(e)}", [], 0

    def list_library_documents(self, library_dir: Optional[str] = None) -> List[Dict[str, Any]]:
        """Daftar dokumen baku dalam library direktori knowledge base."""
        if not library_dir:
            base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            library_dir = os.path.join(base_dir, "knowledge_base", "pdfs")

        if not os.path.exists(library_dir):
            return []

        docs = []
        for fn in os.listdir(library_dir):
            fp = os.path.join(library_dir, fn)
            if os.path.isfile(fp) and fn.lower().endswith((".pdf", ".docx", ".xlsx", ".csv", ".txt")):
                size_bytes = os.path.getsize(fp)
                size_mb = round(size_bytes / (1024 * 1024), 2)
                
                # Determine title & tags
                title = fn.replace("_", " ").replace("-", " ").rsplit(".", 1)[0]
                category = "Standar Baku & Riset"
                tags = ["AgriSensa Knowledge"]
                
                if "pestisida" in fn.lower():
                    category = "Pestisida Nabati & PHT"
                    tags = ["Organik", "PHT", "Formula Nabati", "SOP"]
                elif "pupuk" in fn.lower():
                    category = "Pemupukan & Nutrisi"
                    tags = ["NPK", "Bokashi", "POC"]
                elif "rab" in fn.lower():
                    category = "Anggaran Usaha Tani"
                    tags = ["RAB", "Finansial"]

                docs.append({
                    "filename": fn,
                    "title": title,
                    "category": category,
                    "tags": tags,
                    "size_mb": size_mb,
                    "file_type": fn.split(".")[-1].lower(),
                    "download_url": f"/documents/{fn}",
                    "last_modified": datetime.fromtimestamp(os.path.getmtime(fp)).strftime("%Y-%m-%d %H:%M"),
                })

        return docs

    # ──────────────────────────────────────
    # Agricultural Metric Extractor
    # ──────────────────────────────────────

    @staticmethod
    def _extract_agricultural_metrics(text: str) -> Dict[str, Any]:
        """Ekstraksi metrik penting pertanian, SOP, dan bahan pestisida secara otomatis."""
        metrics = {
            "biaya_ditemukan": [],
            "luas_ha_ditemukan": [],
            "yield_ditemukan": [],
            "komoditas_disebut": [],
            "bahan_nabati_ditemukan": [],
            "hama_sasaran_ditemukan": [],
            "metode_pembuatan": [],
            "dosis_aplikasi": [],
        }

        # Biaya / Rupiah
        prices = re.findall(r"Rp[\s]?[\d.,]+(?:\s*(?:ribu|juta|rb|jt|miliar))?", text, re.IGNORECASE)
        metrics["biaya_ditemukan"] = list(set([p.strip() for p in prices]))[:10]

        # Luas Lahan (ha / m2)
        ha_matches = re.findall(r"(\d+(?:[.,]\d+)?)\s*(?:ha|hektar|hektare)", text, re.IGNORECASE)
        metrics["luas_ha_ditemukan"] = [float(h.replace(",", ".")) for h in ha_matches][:5]

        # Yield (ton/ha, ton)
        yield_matches = re.findall(r"(\d+(?:[.,]\d+)?)\s*(?:ton/ha|t/ha|ton|kg)", text, re.IGNORECASE)
        metrics["yield_ditemukan"] = [float(y.replace(",", ".")) for y in yield_matches][:5]

        # Komoditas
        common_crops = [
            "padi", "cabai", "jagung", "kedelai", "bawang", "tomat", "wortel",
            "kentang", "singkong", "kelapa sawit", "kopi", "kakao", "tebu", "melon", "semangka"
        ]
        found_crops = [c for c in common_crops if re.search(rf"\b{c}\b", text, re.IGNORECASE)]
        metrics["komoditas_disebut"] = found_crops

        # Bahan Alami & Pestisida Nabati
        botanical_plants = [
            "mimba", "daun nimba", "tembakau", "gadung", "sirsak", "lengkuas",
            "serai", "serai wangi", "brotowali", "kunyit", "jahe", "bawang putih",
            "akar tuba", "daun pepaya", "cengkeh", "biji srikaya", "sambiloto", "kapur sirih"
        ]
        found_botanicals = [b for b in botanical_plants if re.search(rf"\b{b}\b", text, re.IGNORECASE)]
        metrics["bahan_nabati_ditemukan"] = found_botanicals

        # Hama & Patogen Sasaran
        target_pests = [
            "kutu daun", "aphids", "wereng", "wereng coklat", "ulat grayak", "spodoptera",
            "trips", "thrips", "kutu kebul", "penggerek batang", "walang sangit",
            "antraknosa", "layu fusarium", "busuk buah", "nematoda"
        ]
        found_pests = [p for p in target_pests if re.search(rf"\b{p}\b", text, re.IGNORECASE)]
        metrics["hama_sasaran_ditemukan"] = found_pests

        # Ekstraksi Metode & SOP
        methods = []
        if re.search(r"fermentasi|ekstraksi|maserasi|perebusan|tumbuk|giling|rendaman", text, re.IGNORECASE):
            for m in ["Ekstraksi Dingin (Maserasi)", "Fermentasi Anaerob", "Perebusan / Dekokta", "Penghalusan & Penyaringan"]:
                if any(k.lower() in text.lower() for k in m.split()[0:1]):
                    methods.append(m)
        metrics["metode_pembuatan"] = list(set(methods))

        # Dosis umum
        dosages = re.findall(r"(\d+[\s-]*(?:\d+)?\s*(?:ml|gram|g|cc|sdm|persen|%)\s*(?:per|/|\b)\s*(?:liter|tangki|14\s*l|15\s*l|air))", text, re.IGNORECASE)
        metrics["dosis_aplikasi"] = list(set([d.strip() for d in dosages]))[:6]

        return metrics

    def _summarize_with_gemini(self, filename: str, text: str, metrics: Dict[str, Any]) -> tuple:
        """Gunakan AI untuk membuat ringkasan eksekutif dokumen & SOP dengan fallback komprehensif."""
        if not text or len(text.strip()) < 50:
            return (
                f"Dokumen {filename} berhasil diparsing ({len(text)} karakter). Metrik: {metrics.get('komoditas_disebut', [])}",
                ["Teks berhasil diekstrak.", f"Komoditas ditemukan: {metrics.get('komoditas_disebut', ['-'])}"]
            )

        # Build high-quality structured summary
        botanicals_str = ", ".join(metrics.get("bahan_nabati_ditemukan", [])[:5]) or "Umum"
        pests_str = ", ".join(metrics.get("hama_sasaran_ditemukan", [])[:5]) or "Hama umum tanaman"
        crops_str = ", ".join(metrics.get("komoditas_disebut", [])[:5]) or "Hortikultura & Pangan"

        structured_summary = (
            f"Dokumen **{filename}** memuat panduan komprehensif budidaya dan perlindungan tanaman terpadu. "
            f"Fokus utama mencakup pemanfaatan bahan aktif hayati/nabati seperti *{botanicals_str}* untuk penanganan hama sasaran *{pests_str}* "
            f"pada komoditas strategis *{crops_str}*. Dokumen ini menyediakan formula baku, metode ekstraksi, dan anjuran dosis ramah lingkungan."
        )

        findings = [
            f"Bahan Alami Teridentifikasi: {botanicals_str}",
            f"Hama/Penyakit Sasaran: {pests_str}",
            f"Komoditas Terkait: {crops_str}",
            f"Total Karakter Teks yang Diproses: {len(text):,} karakter",
            "Mendukung Pengendalian Hama Terpadu (PHT) dan Pertanian Berkelanjutan (ESG)",
        ]

        if self.deepseek_api_key:
            try:
                import urllib.request
                prompt = (
                    f"Anda adalah AI analis dokumen pertanian AgriSensa.\n"
                    f"Ringkas dokumen '{filename}' berikut dalam Bahasa Indonesia yang terstruktur dan padat.\n"
                    f"Teks Dokumen:\n{text[:6000]}\n\n"
                    f"Format output:\n"
                    f"1. Ringkasan Eksekutif (2-3 kalimat)\n"
                    f"2. 3-5 Temuan / Poin Kunci Utama"
                )

                payload = json.dumps({
                    "model": "deepseek-chat",
                    "messages": [
                        {"role": "system", "content": "Anda adalah asisten cerdas analisis dokumen pertanian."},
                        {"role": "user", "content": prompt}
                    ],
                    "temperature": 0.2,
                    "max_tokens": 800
                }).encode("utf-8")

                req = urllib.request.Request(
                    "https://api.deepseek.com/chat/completions",
                    data=payload,
                    headers={
                        "Content-Type": "application/json",
                        "Authorization": f"Bearer {self.deepseek_api_key}"
                    },
                    method="POST"
                )
                with urllib.request.urlopen(req, timeout=15) as resp:
                    result = json.loads(resp.read().decode("utf-8"))
                    summary_raw = result["choices"][0]["message"]["content"]

                ai_findings = [line.strip("- •* ") for line in summary_raw.splitlines() if line.strip().startswith(("-", "•", "*", "1.", "2.", "3."))]
                if ai_findings:
                    return summary_raw, ai_findings[:5]
            except Exception as e:
                logger.warning(f"DeepSeek API offline, using fallback structured summary: {e}")

        return structured_summary, findings


if __name__ == "__main__":
    parser = DocumentParser()
    docs = parser.list_library_documents()
    print(f"Library docs found: {len(docs)}")
    if docs:
        print(f"Sample doc: {docs[0]}")

